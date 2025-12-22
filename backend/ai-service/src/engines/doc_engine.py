"""
Document Engine - 文档生成引擎
生成Word、PDF等文档文件
"""
import os
import json
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import ollama
from openai import AsyncOpenAI
from src.llm.llm_client import get_llm_client


class DocumentEngine:
    """文档生成引擎"""

    def __init__(self):
        self.llm_provider = os.getenv("LLM_PROVIDER", "ollama")
        self.model_name = os.getenv("LLM_MODEL", "qwen2:7b")
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")

        self.client = None
        self.llm_client = None
        self._ready = True

        if self.llm_provider == "openai":
            if self.openai_api_key:
                self.client = AsyncOpenAI(api_key=self.openai_api_key, base_url=self.openai_base_url)
            else:
                self._ready = False
        elif self.llm_provider != "ollama":
            try:
                self.llm_client = get_llm_client()
            except Exception as e:
                print(f"LLM client initialization error: {e}")
                self._ready = False

    def is_ready(self) -> bool:
        """检查引擎是否就绪"""
        return self._ready

    async def generate(
        self,
        prompt: str,
        context: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        生成文档项目

        Args:
            prompt: 用户需求描述
            context: 意图识别结果

        Returns:
            {
                "files": [
                    {"path": "report.docx", "content": bytes, "type": "word"},
                    {"path": "report.pdf", "content": bytes, "type": "pdf"}
                ],
                "metadata": {...}
            }
        """
        if not self._ready:
            raise Exception("Document engine not ready")

        try:
            # 提取上下文信息
            # 防御性编程：确保context是字典而不是列表
            if isinstance(context, list):
                context = context[0] if len(context) > 0 else {}
            entities = context.get("entities", {}) if context else {}
            doc_type = entities.get("doc_type", "report")
            output_format = entities.get("format", "word")  # word, pdf, both

            # 生成文档大纲
            outline = await self._generate_outline(prompt, doc_type, entities)

            # 生成文档内容
            sections = await self._generate_content(outline)

            files = []

            # 生成Word文档
            if output_format in ["word", "both"]:
                word_bytes = self._create_word_document(outline, sections)
                files.append({
                    "path": f"{outline.get('title', 'document')}.docx",
                    "content": word_bytes,
                    "type": "word"
                })

            # 生成PDF文档
            if output_format in ["pdf", "both"]:
                pdf_bytes = self._create_pdf_document(outline, sections)
                files.append({
                    "path": f"{outline.get('title', 'document')}.pdf",
                    "content": pdf_bytes,
                    "type": "pdf"
                })

            return {
                "files": files,
                "metadata": {
                    "doc_type": doc_type,
                    "outline": outline,
                    "format": output_format
                }
            }

        except Exception as e:
            print(f"Document generation error: {e}")
            raise

    async def _generate_outline(
        self,
        prompt: str,
        doc_type: str,
        entities: Dict[str, Any]
    ) -> Dict[str, Any]:
        """生成文档大纲"""

        outline_prompt = f"""根据用户需求生成文档大纲。

用户需求: {prompt}
文档类型: {doc_type}
提取的实体: {json.dumps(entities, ensure_ascii=False)}

请生成包含以下信息的JSON大纲:
{{
    "title": "文档标题",
    "subtitle": "副标题（可选）",
    "author": "作者",
    "date": "2024-01-01",
    "sections": [
        {{
            "title": "第一章 引言",
            "subsections": ["1.1 背景", "1.2 目的"]
        }},
        {{
            "title": "第二章 正文",
            "subsections": ["2.1 概述", "2.2 详细说明"]
        }}
    ],
    "include_toc": true,
    "include_charts": false
}}

只返回JSON，不要其他文字:
"""

        try:
            if self.llm_provider == "ollama":
                response = ollama.chat(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "你是文档结构专家，擅长设计清晰的文档大纲。"},
                        {"role": "user", "content": outline_prompt}
                    ],
                    options={"temperature": 0.3}
                )
                content = response['message']['content']
            elif self.client is not None:
                response = await self.client.chat.completions.create(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "你是文档结构专家，擅长设计清晰的文档大纲。"},
                        {"role": "user", "content": outline_prompt}
                    ],
                    temperature=0.3
                )
                content = response.choices[0].message.content
            elif self.llm_client is not None:
                content = await self.llm_client.chat(
                    messages=[
                        {"role": "system", "content": "You are a document outline assistant. Return valid JSON only."},
                        {"role": "user", "content": outline_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=1024
                )
            else:
                raise Exception("LLM client not initialized")

            outline = json.loads(content)
            return outline

        except Exception as e:
            print(f"Outline generation error: {e}")
            # 返回默认大纲
            return {
                "title": "工作报告",
                "subtitle": "",
                "author": "AI助手",
                "date": "2024-01-01",
                "sections": [
                    {"title": "概述", "subsections": []},
                    {"title": "详细内容", "subsections": []}
                ],
                "include_toc": True,
                "include_charts": False
            }

    async def _generate_content(self, outline: Dict[str, Any]) -> List[Dict[str, Any]]:
        """生成文档各章节内容"""

        sections_content = []

        for section in outline.get("sections", []):
            section_title = section.get("title", "")
            subsections = section.get("subsections", [])

            # 生成章节内容
            content_prompt = f"""为文档章节生成内容。

章节标题: {section_title}
子章节: {", ".join(subsections) if subsections else "无"}

要求:
1. 内容充实、结构清晰
2. 语言专业、简洁
3. 每个段落200-300字
4. 如果有子章节，分别生成内容

直接返回文本内容，不要JSON格式:
"""

            try:
                if self.llm_provider == "ollama":
                    response = ollama.chat(
                        model=self.model_name,
                        messages=[
                            {"role": "system", "content": "你是专业的文档撰写专家。"},
                            {"role": "user", "content": content_prompt}
                        ],
                        options={"temperature": 0.5, "num_predict": 1024}
                    )
                    content = response['message']['content']
                elif self.client is not None:
                    response = await self.client.chat.completions.create(
                        model=self.model_name,
                        messages=[
                            {"role": "system", "content": "你是专业的文档撰写专家。"},
                            {"role": "user", "content": content_prompt}
                        ],
                        temperature=0.5,
                        max_tokens=1024
                    )
                    content = response.choices[0].message.content
                elif self.llm_client is not None:
                    content = await self.llm_client.chat(
                        messages=[
                            {"role": "system", "content": "You are a document writer. Return plain text."},
                            {"role": "user", "content": content_prompt}
                        ],
                        temperature=0.5,
                        max_tokens=1024
                    )
                else:
                    raise Exception("LLM client not initialized")

                sections_content.append({
                    "title": section_title,
                    "content": content.strip(),
                    "subsections": subsections
                })

            except Exception as e:
                print(f"Content generation error for section '{section_title}': {e}")
                sections_content.append({
                    "title": section_title,
                    "content": f"这是{section_title}的内容。",
                    "subsections": subsections
                })

        return sections_content

    def _create_word_document(
        self,
        outline: Dict[str, Any],
        sections: List[Dict[str, Any]]
    ) -> bytes:
        """创建Word文档"""

        doc = Document()

        # 设置标题
        title = doc.add_heading(outline.get("title", "文档"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题
        if outline.get("subtitle"):
            subtitle = doc.add_heading(outline.get("subtitle", ""), 2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 作者和日期
        author_para = doc.add_paragraph()
        author_para.add_run(f"作者: {outline.get('author', 'AI助手')}\n").italic = True
        author_para.add_run(f"日期: {outline.get('date', '2024-01-01')}").italic = True
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # 目录（简化版）
        if outline.get("include_toc", True):
            doc.add_heading("目录", 1)
            for idx, section in enumerate(sections, 1):
                doc.add_paragraph(f"{idx}. {section['title']}", style='List Number')
            doc.add_page_break()

        # 各章节内容
        for section in sections:
            doc.add_heading(section["title"], 1)
            doc.add_paragraph(section["content"])
            doc.add_paragraph()  # 空行

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _create_pdf_document(
        self,
        outline: Dict[str, Any],
        sections: List[Dict[str, Any]]
    ) -> bytes:
        """创建PDF文档"""

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []

        styles = getSampleStyleSheet()

        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=1  # 居中
        )

        # 章节标题样式
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=12
        )

        # 正文样式
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=16,
            spaceAfter=12
        )

        # 添加标题
        story.append(Paragraph(outline.get("title", "文档"), title_style))

        # 副标题
        if outline.get("subtitle"):
            story.append(Paragraph(outline.get("subtitle", ""), styles['Heading2']))

        # 作者和日期
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"作者: {outline.get('author', 'AI助手')}", styles['Italic']))
        story.append(Paragraph(f"日期: {outline.get('date', '2024-01-01')}", styles['Italic']))

        story.append(PageBreak())

        # 目录
        if outline.get("include_toc", True):
            story.append(Paragraph("目录", heading_style))
            for idx, section in enumerate(sections, 1):
                story.append(Paragraph(f"{idx}. {section['title']}", body_style))
            story.append(PageBreak())

        # 各章节内容
        for section in sections:
            story.append(Paragraph(section["title"], heading_style))
            story.append(Paragraph(section["content"], body_style))
            story.append(Spacer(1, 0.3*inch))

        # 生成PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
