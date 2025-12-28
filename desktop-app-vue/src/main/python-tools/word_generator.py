#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档生成工具
使用python-docx库创建和编辑Word文档
"""

import sys
import json
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(json.dumps({
        'success': False,
        'error': '未安装python-docx，请运行: pip install python-docx'
    }, ensure_ascii=False))
    sys.exit(1)

class WordGenerator:
    """Word文档生成器"""

    def __init__(self):
        self.doc = Document()

    def create_document(self, params):
        """
        创建Word文档

        参数:
        - title: 文档标题
        - content: 文档内容（可以是字符串或段落列表）
        - output_path: 输出文件路径
        - template: 模板类型 (business/academic/report)
        - metadata: 文档元数据（作者、主题等）
        """
        title = params.get('title', '未命名文档')
        content = params.get('content', '')
        output_path = params.get('output_path')
        template = params.get('template', 'basic')
        metadata = params.get('metadata', {})

        # 设置文档元数据
        if metadata:
            if 'author' in metadata:
                self.doc.core_properties.author = metadata['author']
            if 'subject' in metadata:
                self.doc.core_properties.subject = metadata['subject']
            if 'keywords' in metadata:
                self.doc.core_properties.keywords = metadata['keywords']

        # 根据模板类型设置样式
        if template == 'business':
            self._apply_business_template(title, content)
        elif template == 'academic':
            self._apply_academic_template(title, content)
        elif template == 'report':
            self._apply_report_template(title, content)
        else:
            self._apply_basic_template(title, content)

        # 保存文档
        if not output_path:
            output_path = f"{title}.docx"

        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        self.doc.save(output_path)

        return {
            'success': True,
            'output_path': output_path,
            'file_size': os.path.getsize(output_path),
            'created_at': datetime.now().isoformat()
        }

    def _apply_basic_template(self, title, content):
        """基础模板"""
        # 添加标题
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加内容
        if isinstance(content, list):
            for para in content:
                self.doc.add_paragraph(para)
        else:
            self.doc.add_paragraph(content)

    def _apply_business_template(self, title, content):
        """商务模板"""
        # 标题（加粗、蓝色）
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 153)  # 深蓝色

        # 日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"日期: {datetime.now().strftime('%Y年%m月%d日')}")
        date_run.font.size = Pt(10)
        date_run.font.italic = True

        self.doc.add_paragraph()  # 空行

        # 内容
        if isinstance(content, str):
            content = content.split('\n\n')

        for para_text in content:
            if para_text.strip():
                para = self.doc.add_paragraph(para_text.strip())
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_after = Pt(12)

    def _apply_academic_template(self, title, content):
        """学术模板"""
        # 标题
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 摘要（如果content是字典且包含abstract）
        if isinstance(content, dict):
            if 'abstract' in content:
                abstract_heading = self.doc.add_heading('摘要', 1)
                self.doc.add_paragraph(content['abstract'])
                self.doc.add_paragraph()

            if 'body' in content:
                content = content['body']

        # 正文
        if isinstance(content, list):
            for para in content:
                self.doc.add_paragraph(para)
        else:
            self.doc.add_paragraph(content)

        # 参考文献（示例）
        ref_heading = self.doc.add_heading('参考文献', 1)
        self.doc.add_paragraph('（请在此添加参考文献）', style='List Number')

    def _apply_report_template(self, title, content):
        """报告模板"""
        # 封面
        cover_title = self.doc.add_heading(title, 0)
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cover_title.runs:
            run.font.size = Pt(28)
            run.font.bold = True

        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run('工作报告')
        subtitle_run.font.size = Pt(18)

        # 日期
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime('%Y年%m月%d日'))
        date_run.font.size = Pt(14)

        # 分页
        self.doc.add_page_break()

        # 目录（占位）
        toc_heading = self.doc.add_heading('目录', 1)
        self.doc.add_paragraph('（请在此插入目录）')
        self.doc.add_page_break()

        # 正文
        if isinstance(content, dict):
            # 按章节处理
            for section_title, section_content in content.items():
                self.doc.add_heading(section_title, 1)
                if isinstance(section_content, list):
                    for para in section_content:
                        self.doc.add_paragraph(para)
                else:
                    self.doc.add_paragraph(section_content)
                self.doc.add_paragraph()
        elif isinstance(content, list):
            for para in content:
                self.doc.add_paragraph(para)
        else:
            self.doc.add_paragraph(content)

def main():
    """主函数"""
    try:
        # 解析参数
        if len(sys.argv) < 2:
            raise ValueError('缺少参数')

        args = json.loads(sys.argv[1])

        # 验证必需参数
        if 'operation' not in args:
            args['operation'] = 'create'

        generator = WordGenerator()

        # 执行操作
        if args['operation'] == 'create':
            result = generator.create_document(args)
        else:
            raise ValueError(f"不支持的操作: {args['operation']}")

        # 输出结果
        print(json.dumps(result, ensure_ascii=False, indent=2))

    except Exception as e:
        error_result = {
            'success': False,
            'error': str(e),
            'type': type(e).__name__
        }
        print(json.dumps(error_result, ensure_ascii=False))
        sys.exit(1)

if __name__ == '__main__':
    main()
